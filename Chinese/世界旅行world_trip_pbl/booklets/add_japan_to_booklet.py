#!/usr/bin/env python3
"""Add Japan section to the existing asia_booklet_AwithChina.docx"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = '/Users/Huan/projects/summercourse/Chinese/世界旅行world_trip_pbl/booklets/asia_booklet_AwithChina.docx'
OUT = '/Users/Huan/projects/summercourse/Chinese/世界旅行world_trip_pbl/booklets/asia_booklet_AwithChina.docx'

doc = Document(SRC)
body = doc.element.body

JAPAN_RED = "BC002D"

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)
    return p

def add_heading_para(doc, text, size=Pt(28), color=JAPAN_RED, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_shaded_bar(doc, text, fill_color=JAPAN_RED):
    p = doc.add_paragraph()
    # Add shading to paragraph
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run("  " + text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return p

def add_text(doc, text, size=None, bold=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if size:
        run.font.size = size
    if bold:
        run.font.bold = bold
    return p

def add_empty_line(doc):
    doc.add_paragraph()

# ═══════════════════════════════════════
# Add page break before Japan section
# ═══════════════════════════════════════
add_page_break(doc)

# ═══════════════════════════════════════
# JAPAN PAGE - following China's structure
# ═══════════════════════════════════════

# Title
add_heading_para(doc, "🇯🇵 日本 Japan", size=Pt(28), color=JAPAN_RED)

# 涂一涂国旗
add_shaded_bar(doc, "涂一涂: 日本国旗 Color the Flag")
add_text(doc, "⬜ 白色 White + 🔴 红色 Red（中间圆形）")
add_empty_line(doc)  # space for flag drawing

# 描一描
add_shaded_bar(doc, "描一描 Trace")
# Leave space for tracing 日本 characters
p = doc.add_paragraph()
run = p.add_run("      ")  # placeholder space for trace boxes
add_empty_line(doc)

# 圈一圈
add_shaded_bar(doc, "⭕ 圈一圈 Circle the Correct Answer")

add_text(doc, "1. 在日本见面应该怎么做？ How do you greet in Japan?")
add_text(doc, "   A. 拥抱 Hug     B. 握手 Shake hands     C. 鞠躬 Bow")
add_empty_line(doc)

add_text(doc, "2. 进别人家要做什么？ What do you do when entering a house?  ")
add_text(doc, " A. 脱鞋 Take off shoes     B. 洗手 Wash hands     C. 拍手 Clap")
add_empty_line(doc)

add_text(doc, "3. 在日本吃拉面发出声音是？ Slurping ramen in Japan is:  ")
add_text(doc, " A. 不礼貌 Rude     B. 礼貌 Polite     C. 奇怪 Weird")

# 在地图中找出日本
add_shaded_bar(doc, "在地图中找出日本，并涂色标注。 Find Japan on the map, color it, and label it.")
add_empty_line(doc)  # space for map

# 画一画写一写
add_shaded_bar(doc, "画一画,写一写。Draw & Write. ")
p = doc.add_paragraph()
run = p.add_run("我最喜欢的日本食物 My Favorite Japanese Food")
run.font.size = Pt(12)
run.font.bold = True

# Add a table for drawing area (like China section)
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.text = "比如：寿司、拉面、天妇罗。For example: sushi, ramen, tempura."
for _ in range(4):
    cell.add_paragraph("")

add_empty_line(doc)

# Save
doc.save(OUT)
print(f"Japan section added. Saved to {OUT}")
