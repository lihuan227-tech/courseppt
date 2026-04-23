#!/usr/bin/env python3
"""Convert asia_booklet_AwithChina.docx to HTML with embedded images."""
import base64
import html as html_mod
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = '/Users/Huan/projects/summercourse/Chinese/世界旅行world_trip_pbl/booklets/asia_booklet_AwithChina.docx'
OUT = '/Users/Huan/projects/summercourse/Chinese/世界旅行world_trip_pbl/booklets/asia_booklet_AwithChina.html'

doc = Document(SRC)

# Build rId -> base64 image map
img_map = {}
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        blob = rel.target_part.blob
        ct = rel.target_part.content_type
        b64 = base64.b64encode(blob).decode('ascii')
        img_map[rel.rId] = f"data:{ct};base64,{b64}"

def get_image_from_paragraph(para_element):
    """Extract image rId from a paragraph's drawing/inline elements."""
    images = []
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    # Find all blip elements (inline images)
    for blip in para_element.iter(qn('a:blip')):
        rId = blip.get(qn('r:embed'))
        if rId and rId in img_map:
            # Get dimensions
            parent_inline = None
            for ancestor in blip.iterancestors():
                if ancestor.tag == qn('wp:inline') or ancestor.tag == qn('wp:anchor'):
                    parent_inline = ancestor
                    break
            w, h = 300, 200  # defaults
            if parent_inline is not None:
                extent = parent_inline.find(qn('wp:extent'))
                if extent is not None:
                    emu_w = int(extent.get('cx', 0))
                    emu_h = int(extent.get('cy', 0))
                    if emu_w > 0:
                        w = int(emu_w / 9525)  # EMU to px
                        h = int(emu_h / 9525)
            images.append((rId, w, h))
    return images

# Color mapping
SHADING_COLORS = {
    'DE2910': 'shaded',       # China red
    'BC002D': 'shaded japan',  # Japan red
    'FF8F00': 'shaded accent', # Orange accent
    '138808': 'shaded india',  # India green
}

out = []
out.append('''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>探索亚洲 Explore Asia — Global Explorer Camp</title>
<style>
  body { font-family: "KaiTi", "Noto Sans SC", "Kaiti SC", sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; color: #333; line-height: 1.6; }
  .shaded { background: #DE2910; color: white; padding: 6px 12px; font-weight: bold; font-size: 14px; margin: 12px 0 6px; border-radius: 4px; }
  .shaded.japan { background: #BC002D; }
  .shaded.india { background: #138808; }
  .shaded.accent { background: #FF8F00; }
  .heading { font-size: 28px; font-weight: bold; margin: 20px 0 10px; }
  .heading.china { color: #DE2910; }
  .heading.japan { color: #BC002D; }
  .heading.india { color: #138808; }
  .heading.asia { color: #C62828; }
  table { border-collapse: collapse; width: 100%; margin: 10px 0; }
  td, th { border: 1px solid #ccc; padding: 8px; text-align: center; vertical-align: middle; }
  th { background: #FF8F00; color: white; }
  img { max-width: 100%; height: auto; display: block; margin: 8px auto; }
  .img-inline { display: inline-block; vertical-align: middle; margin: 4px; }
  .page-break { page-break-after: always; border-top: 3px solid #eee; margin: 40px 0; padding-top: 10px; }
  .question { margin: 8px 0 2px; }
  .options { margin: 0 0 8px 20px; color: #555; }
  .draw-area { border: 2px dashed #bbb; min-height: 150px; border-radius: 10px; margin: 10px 0; padding: 10px; }
  @media print {
    .page-break { border: none; page-break-after: always; }
    body { max-width: 100%; }
  }
</style>
</head>
<body>
''')

# Process each element in document body
body_elements = list(doc.element.body)
para_index = 0
table_index = 0
all_paras = list(doc.paragraphs)
all_tables = list(doc.tables)

for element in body_elements:
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

    if tag == 'p':
        # Find matching paragraph
        para = None
        for p in all_paras:
            if p._element is element:
                para = p
                break
        if para is None:
            continue

        text = para.text.strip()

        # Check for images in this paragraph
        images = get_image_from_paragraph(element)

        # Check for page break
        has_pb = False
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_pb = True
        if has_pb:
            out.append('<div class="page-break"></div>')
            continue

        # Check shading
        pPr = element.find(qn('w:pPr'))
        shading = None
        if pPr is not None:
            shd = pPr.find(qn('w:shd'))
            if shd is not None:
                shading = shd.get(qn('w:fill'))

        # Check font properties
        font_size = None
        is_bold = False
        font_color = None
        for run in para.runs:
            if run.font.size and (font_size is None or run.font.size > font_size):
                font_size = run.font.size
            if run.font.bold:
                is_bold = True
            try:
                if run.font.color.rgb:
                    font_color = str(run.font.color.rgb)
            except:
                pass

        escaped = html_mod.escape(text)

        # Output images first if any
        if images:
            for rId, w, h in images:
                src = img_map.get(rId, '')
                if src:
                    max_w = min(w, 700)
                    out.append(f'<img src="{src}" width="{max_w}" style="max-width:100%;">')

        if not text and not images:
            out.append('<br>')
            continue

        if not text and images:
            continue  # image already output

        if shading:
            css = SHADING_COLORS.get(shading, 'shaded')
            out.append(f'<div class="{css}">{escaped}</div>')
        elif font_size and font_size >= 300000:
            css = 'heading'
            if '中国' in text and 'China' in text:
                css += ' china'
            elif '日本' in text and 'Japan' in text:
                css += ' japan'
            elif '印度' in text and 'India' in text:
                css += ' india'
            elif '亚洲' in text or '🌏' in text:
                css += ' asia'
            out.append(f'<div class="{css}">{escaped}</div>')
        elif is_bold and font_size and font_size >= 150000:
            out.append(f'<p><strong>{escaped}</strong></p>')
        else:
            out.append(f'<p>{escaped}</p>')

    elif tag == 'tbl':
        for table in all_tables:
            if table._element is element:
                out.append('<table>')
                for ri, row in enumerate(table.rows):
                    out.append('<tr>')
                    for ci, cell in enumerate(row.cells):
                        cell_text = html_mod.escape(cell.text.strip())
                        # Check for images in table cells
                        cell_images = []
                        for cp in cell.paragraphs:
                            cell_images.extend(get_image_from_paragraph(cp._element))

                        tag_name = 'th' if ri == 0 and len(table.rows) > 2 else 'td'
                        cell_html = cell_text
                        for rId, w, h in cell_images:
                            src = img_map.get(rId, '')
                            if src:
                                max_w = min(w, 250)
                                cell_html += f'<br><img src="{src}" width="{max_w}" style="max-width:100%;">'
                        out.append(f'<{tag_name}>{cell_html}</{tag_name}>')
                    out.append('</tr>')
                out.append('</table>')
                break

# Now wrap content into pages by splitting on page-break divs
html_content = '\n'.join(out)

# Split into pages
import re
# Find all content, split by page-break markers
parts = html_content.split('<div class="page-break"></div>')

final = []
# Everything before first page break is already output with headers
# We need to re-wrap: take the header (everything up to <body> + nav), then wrap each part in <div class="page">

# Find the end of <body> + nav section
body_start = html_content.find('<body>')
header_html = html_content[:html_content.find('\n', body_start) + 1]

# Actually, let's just rebuild from parts
final_out = []
final_out.append('''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>探索亚洲 Explore Asia — Global Explorer Camp</title>
<style>
  @page { size: letter; margin: 0.5in; }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: "KaiTi", "Noto Sans SC", "Kaiti SC", sans-serif; color: #333; line-height: 1.6; background: #f0f0f0; overflow: hidden; height: 100vh; }
  .page { display: none; width: 8.5in; height: 11in; margin: 0 auto; padding: 0.5in; background: #fff; border-radius: 8px; box-shadow: 0 2px 12px rgba(0,0,0,0.15); overflow-y: auto; position: relative; }
  .page.active { display: block; }
  .shaded { background: #DE2910; color: white; padding: 6px 12px; font-weight: bold; font-size: 14px; margin: 12px 0 6px; border-radius: 4px; }
  .shaded.japan { background: #BC002D; }
  .shaded.india { background: #138808; }
  .shaded.accent { background: #FF8F00; }
  .heading { font-size: 28px; font-weight: bold; margin: 20px 0 10px; }
  .heading.china { color: #DE2910; }
  .heading.japan { color: #BC002D; }
  .heading.india { color: #138808; }
  .heading.asia { color: #C62828; }
  table { border-collapse: collapse; width: 100%; margin: 10px 0; }
  td, th { border: 1px solid #ccc; padding: 8px; text-align: center; vertical-align: middle; }
  th { background: #FF8F00; color: white; }
  img { max-width: 100%; height: auto; display: block; margin: 8px auto; }
  .nav { position: fixed; bottom: 0; left: 0; right: 0; background: #333; color: #fff; display: flex; align-items: center; justify-content: center; gap: 30px; padding: 12px; z-index: 999; }
  .nav button { background: #FF8F00; color: #fff; border: none; padding: 10px 24px; border-radius: 6px; font-size: 18px; cursor: pointer; font-family: inherit; }
  .nav button:hover { background: #E67E00; }
  .nav button:disabled { background: #999; cursor: not-allowed; }
  .nav .page-info { font-size: 16px; min-width: 100px; text-align: center; }
  @media screen {
    .page { margin-top: 10px; margin-bottom: 70px; }
  }
  @media print {
    .nav { display: none; }
    .page { display: block !important; page-break-after: always; box-shadow: none; margin: 0; border-radius: 0; overflow: visible; width: 100%; height: auto; min-height: 10in; }
    body { background: #fff; overflow: visible; height: auto; }
  }
</style>
</head>
<body>

<div class="nav">
  <button id="prevBtn" onclick="navigate(-1)">&larr; 上一页</button>
  <span class="page-info" id="pageInfo">1 / ?</span>
  <button id="nextBtn" onclick="navigate(1)">下一页 &rarr;</button>
</div>
''')

# Extract just the content lines (skip header/footer we already wrote)
# Find content between <body> and </body> in original
body_match_start = html_content.find('\n<br>')  # first content line
body_match_end = html_content.rfind('</body>')
content_html = html_content[body_match_start:body_match_end].strip()

# Split by page-break divs
page_contents = content_html.split('<div class="page-break"></div>')

for i, page_html in enumerate(page_contents):
    page_html = page_html.strip()
    if not page_html:
        continue
    active = ' active' if i == 0 else ''
    final_out.append(f'<div class="page{active}">')
    final_out.append(page_html)
    final_out.append('</div>')

# Add navigation JS
final_out.append('''
<script>
let current = 0;
const pages = document.querySelectorAll('.page');
const total = pages.length;

function navigate(dir) {
  pages[current].classList.remove('active');
  current = Math.max(0, Math.min(total - 1, current + dir));
  pages[current].classList.add('active');
  pages[current].scrollTop = 0;
  update();
}

function update() {
  document.getElementById('pageInfo').textContent = (current + 1) + ' / ' + total;
  document.getElementById('prevBtn').disabled = current === 0;
  document.getElementById('nextBtn').disabled = current === total - 1;
}

document.addEventListener('keydown', e => {
  if (e.key === 'ArrowRight' || e.key === 'ArrowDown') { navigate(1); e.preventDefault(); }
  if (e.key === 'ArrowLeft' || e.key === 'ArrowUp') { navigate(-1); e.preventDefault(); }
});

update();
</script>
''')

final_out.append('</body></html>')

with open(OUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(final_out))

print(f"Created paginated HTML with {len(img_map)} images, {len(page_contents)} pages → {OUT}")
