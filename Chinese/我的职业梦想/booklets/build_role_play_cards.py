#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Day 4 · 小小职业剧场 情景卡 — printable cut-out cards (Word version).

Generates `role_play_cards.docx` — 12 scenario cards (4 teacher · 4 doctor · 4 chef)
formatted as a 2x2 grid per page so the teacher can print and cut them out.

Big emojis are used as visual cues so children who don't yet read Chinese can
still recognize the situation; CN + EN names + a short prompt accompany each.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Card data (mirrors create_day4_helpers.py) -------------------------------
PROFESSIONS = [
    {
        "label_cn": "👩‍🏫 老师 Teacher",
        "color":    "43A047",   # teacher green
        "role":     "1 人 演 老师 · 1-2 人 演 学生 (其他 是 观众)",
        "cards": [
            ("🗣️", "学生 一直 讲话", "Students won't stop talking",
             "你 是 老师 — 学生 一直 在 聊天, 你 怎么 让 大家 安静?"),
            ("😢", "有人 哭 了", "Someone is crying",
             "你 是 老师 — 一个 学生 哭 了, 你 怎么 安慰 他?"),
            ("🤔", "有 人 不会 做 题", "Someone can't do the work",
             "你 是 老师 — 一个 学生 说「我 不会」, 你 怎么 帮 他?"),
            ("🏃", "大家 不会 排队", "No one is lining up",
             "你 是 老师 — 大家 跑 来 跑 去, 你 怎么 让 大家 排队?"),
        ],
    },
    {
        "label_cn": "👩‍⚕️ 医生 Doctor",
        "color":    "C8253E",
        "role":     "1 人 演 医生 · 1-2 人 演 病人 (其他 是 观众)",
        "cards": [
            ("🤢", "肚子 疼", "Stomach ache",
             "你 是 医生 — 一个 小朋友 说 肚子 疼, 你 怎么 帮 他?"),
            ("🤕", "摔倒 哭 了", "Fell and crying",
             "你 是 医生 — 一个 小朋友 摔倒 在 哭, 你 怎么 看 他?"),
            ("😨", "害怕 打 针", "Scared of shots",
             "你 是 医生 — 一个 小朋友 怕 打针, 你 怎么 安慰 他?"),
            ("🪑", "很多 病人 排队", "Many patients waiting",
             "你 是 医生 — 候诊室 排了 好多 人, 你 怎么 一个 一个 看?"),
        ],
    },
    {
        "label_cn": "👨‍🍳 厨师 Chef",
        "color":    "FF8F00",
        "role":     "1-2 人 演 厨师 · 1-2 人 演 客人 (其他 是 观众)",
        "cards": [
            ("🤤", "客人 饿 了", "Customer is hungry",
             "你 是 厨师 — 客人 进 餐厅, 你 怎么 招呼 + 给 他 吃 的?"),
            ("🥲", "食物 掉 地上 了", "Food fell down",
             "你 是 厨师 — 你 端 菜 不小心 摔 了, 你 怎么 办?"),
            ("🥵", "客人 说 太 辣", "Too spicy!",
             "你 是 厨师 — 客人 说 你 做 的 菜 太 辣, 你 怎么 办?"),
            ("🍴", "餐厅 太 忙", "Restaurant too busy",
             "你 是 厨师 — 一下 来 好多 客人, 你 怎么 不 慌?"),
        ],
    },
]

# ---- Helpers ----------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, hex_color, sz=16):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), hex_color)
        borders.append(b)
    tc_pr.append(borders)

def add_run(p, text, *, size=14, bold=False, color=None, font="KaiTi"):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r

def add_centered(cell, text, *, size=14, bold=False, color=None, space_before=0, space_after=0):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, size=size, bold=bold, color=color)
    return p

def fill_card(cell, card_num, emoji, cn, en, desc, color, label):
    """Render one scenario card inside a table cell."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_borders(cell, color, sz=24)
    # Clear default paragraph
    cell.paragraphs[0].text = ""
    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p_num, f"  No. {card_num}   ", size=11, bold=True, color="FFFFFF")
    # Set badge background on the run requires shading paragraph — simplify: just colored bold text
    # Profession label small at top
    add_centered(cell, label, size=10, color=color, space_before=2, space_after=2)
    # Big emoji
    add_centered(cell, emoji, size=72, space_before=4, space_after=4)
    # Scenario CN
    add_centered(cell, cn, size=18, bold=True, color=color, space_after=2)
    # Scenario EN
    add_centered(cell, en, size=10, color="888888", space_after=6)
    # Description
    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(2)
    add_run(p_desc, desc, size=11, color="2C2C2C")

# ---- Build document ---------------------------------------------------------
doc = Document()

# Letter, narrow margins
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(0.4)
section.bottom_margin = Inches(0.4)
section.left_margin = Inches(0.4)
section.right_margin = Inches(0.4)

# Cover/intro
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title, "🎭 小小职业剧场 · 情景卡", size=26, bold=True, color="C45A2A")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, "Mini Career Theater · Scenario Cards (printable cut-outs)", size=12, color="888888")
how = doc.add_paragraph()
how.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(how, "✂️ 打印 → 沿 边 剪开 → 每组 抽 一张 → 演 出来 (用 动作 + 对话, 不 说「我 是 ___!」)",
        size=11, bold=True, color="2C2C2C")
how_en = doc.add_paragraph()
how_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(how_en, "Print → cut along the edges → each group draws one → act it out (no spoilers!)",
        size=9, color="888888")
doc.add_paragraph()  # spacer

# One section per profession: heading + 2x2 grid of cards on a single page
for idx, prof in enumerate(PROFESSIONS):
    if idx > 0:
        doc.add_page_break()
    # Heading
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(h, prof["label_cn"] + "  情景卡", size=22, bold=True, color=prof["color"])
    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(role, "👥 分角色: " + prof["role"], size=11, bold=True, color="2C2C2C")
    doc.add_paragraph()  # spacer

    # 2x2 grid of cards
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set column widths
    for row in table.rows:
        row.height = Inches(4.2)
        for cell in row.cells:
            cell.width = Inches(3.85)
    for i, card in enumerate(prof["cards"]):
        r, c = i // 2, i % 2
        cell = table.rows[r].cells[c]
        emoji, cn, en, desc = card
        fill_card(cell, i + 1, emoji, cn, en, desc, prof["color"], prof["label_cn"])

# Save
out_path = os.path.join(os.path.dirname(__file__), "role_play_cards.docx")
doc.save(out_path)
print(f"Saved {out_path}  (12 scenario cards · 3 pages · 4 cards/page)")
